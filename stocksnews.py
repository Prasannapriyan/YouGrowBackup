import time
import undetected_chromedriver as uc
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC

# --- Functions for Word Document Formatting ---
# These are needed for the title's background shading.
try:
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    import docx.shared

    def add_shading_to_paragraph(paragraph, color="000000"):
        """Applies a background color shading to an entire paragraph."""
        shading_xml = f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color}"/>'
        pPr = paragraph._p.get_or_add_pPr()
        pPr.append(parse_xml(shading_xml))

except ImportError:
    print("Warning: python-docx Oxml features not available. Title background may not work.")
    def add_shading_to_paragraph(paragraph, color): pass


def create_stocks_bulletin(output_filename="Key_Stocks_to_Watch.docx"):
    """
    Scrapes market news from Groww's stocks section,
    formats the top 10 items with company details, and saves them to a .docx file.

    Args:
        output_filename (str): The name of the output .docx file.

    Returns:
        bool: True if successful, False otherwise.
    """
    print("Fetching latest stocks news from Groww...")
    
    driver = None
    try:
        # Initialize browser
        print("Initializing browser...")
        options = uc.ChromeOptions()
        options.add_argument('--headless')
        options.add_argument('--no-sandbox')
        options.add_argument('--disable-dev-shm-usage')
        
        # Ensure compatibility by using the correct ChromeDriver version
        driver = uc.Chrome(version_main=137, options=options)
        wait = WebDriverWait(driver, 20)
        
        # Load page
        url = "https://groww.in/market-news/stocks"
        print("Loading page...")
        driver.get(url)
        
        # Wait for initial items to load
        wait.until(EC.presence_of_element_located((By.CLASS_NAME, 'smnli671ItemContainer')))
        
        # Scroll to load all items
        print("Loading more news items...")
        SCROLL_PAUSE_TIME = 2
        scroll_count = 0
        max_scrolls = 2  # Ensure we get enough items
        
        while scroll_count < max_scrolls:
            # Scroll down
            driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
            time.sleep(SCROLL_PAUSE_TIME)
            scroll_count += 1
            print(f"Scroll {scroll_count}/{max_scrolls}")
            
            # Try to click "Load More" button if it exists
            try:
                load_more = driver.find_element(By.CLASS_NAME, "sm403InfiniteLoaderContainer")
                if load_more.is_displayed():
                    driver.execute_script("arguments[0].scrollIntoView(true);", load_more)
                    time.sleep(1)
            except:
                pass
        
        # Get all loaded items
        soup = BeautifulSoup(driver.page_source, 'html.parser')
        news_items = soup.find_all('div', {'class': lambda x: x and 'ItemContainer' in x})
        
        if not news_items:
            print("Could not find news items.")
            return False
            
        print(f"Found {len(news_items)} news items to process")

        print(f"Filtering news and creating '{output_filename}'...")
        document = Document()
        
        # --- Add Title ---
        title_paragraph = document.add_paragraph()
        title_run = title_paragraph.add_run(" Key Stocks to Watch ")
        font = title_run.font
        font.name = 'Arial Black'; font.size = Pt(20)
        if docx.shared: font.color.rgb = docx.shared.RGBColor(0xff, 0xff, 0xff)
        add_shading_to_paragraph(title_paragraph, color="000000")
        document.add_paragraph()

        # --- Process News Items ---
        seen_stocks = set()  # Track unique stock names
        news_count = 0  # Track number of unique news added
        skipped_count = 0  # Track duplicates
        
        for idx, item in enumerate(news_items, 1):
            print(f"\rProcessing news {idx}/{len(news_items)}...", end="")
            if news_count >= 15:  # Get 15 unique items
                break
            try:
                # Extract news details
                # Get the header container
                header_div = item.find('div', class_='smnli671BoxHeaderText')
                if not header_div:
                    header_div = item.find('div', {'class': lambda x: x and 'BoxHeaderText' in x})
                
                if header_div:
                    source = header_div.find('div').text.strip()
                    time_element = header_div.find('time')
                    timestamp = time_element.text.strip() if time_element else ""
                else:
                    print("Warning: Could not find header for news item", news_count + 1)
                    source = "Unknown Source"
                    timestamp = ""
                
                # Get the headline
                headline_div = item.find('div', class_='smnli671BoxItemTitle')
                if not headline_div:
                    headline_div = item.find('div', {'class': lambda x: x and 'BoxItemTitle' in x})
                
                if headline_div:
                    headline = headline_div.text.strip()
                else:
                    print("Warning: Could not find headline for news item", news_count + 1)
                    continue
                
                # Find stock details
                # Find stock details with flexible class matching
                stock_container = item.find('span', class_='smnli671MarketNewsCompName') or \
                                item.find('span', {'class': lambda x: x and 'MarketNewsCompName' in x})
                if stock_container:
                    stock_name = stock_container.text.strip()
                    if stock_name in seen_stocks:
                        skipped_count += 1
                        print(f"  > Skipping duplicate stock: {stock_name}")
                        continue
                    seen_stocks.add(stock_name)
                    
                    # Get price change span element
                    change_element = stock_container.find_next_sibling('span')
                    
                    # Get the class and text
                    if change_element:
                        classes = change_element.get('class', [])
                        price_text = change_element.text.strip()
                        
                        # Determine if positive or negative
                        if 'contentPositive' in classes:
                            price_change = f"{price_text}"
                        elif 'contentNegative' in classes:
                            price_change = f"{price_text}"
                        else:
                            price_change = price_text
                    else:
                        price_change = "N/A"
                else:
                    stock_name = ""
                    price_change = ""
                
                news_count += 1  # Increment counter for unique news
                
                # Add numbered headline
                p_headline = document.add_paragraph()
                p_headline.paragraph_format.left_indent = Inches(0.25)
                p_headline.paragraph_format.first_line_indent = Inches(-0.25)
                p_headline.paragraph_format.space_before = Pt(12)
                p_headline.paragraph_format.space_after = Pt(6)

                # Add number and headline
                number_run = p_headline.add_run(f"{news_count}. ")
                number_run.bold = True
                number_run.font.size = Pt(12)
                
                runner = p_headline.add_run(f"{headline}")
                runner.bold = True
                runner.font.size = Pt(12)
                
                # Add source and time with proper spacing
                p_source = document.add_paragraph()
                p_source.paragraph_format.left_indent = Inches(0.5)
                p_source.paragraph_format.space_before = Pt(3)
                p_source.paragraph_format.space_after = Pt(3)
                source_run = p_source.add_run(f"{source} • {timestamp}")
                source_run.italic = True
                source_run.font.size = Pt(10)
                source_run.font.color.rgb = docx.shared.RGBColor(89, 89, 89)
                
                # Add stock details if available
                if stock_name:
                    p_stock = document.add_paragraph()
                    p_stock.paragraph_format.left_indent = Inches(0.5)
                    p_stock.paragraph_format.space_before = Pt(3)
                    p_stock.paragraph_format.space_after = Pt(12)
                    stock_run = p_stock.add_run(f"{stock_name} ({price_change})")
                    stock_run.bold = True
                    stock_run.font.size = Pt(11)
                    if price_change.startswith('+'):
                        stock_run.font.color.rgb = docx.shared.RGBColor(0, 128, 0)  # Green
                    elif price_change.startswith('-'):
                        stock_run.font.color.rgb = docx.shared.RGBColor(255, 0, 0)  # Red
                else:
                    # Add extra space after source if no stock details
                    p_source.paragraph_format.space_after = Pt(12)
                
            except Exception as e:
                print(f"Error processing news item {news_count + 1}: {e}")
                continue

        document.save(output_filename)
        print(f"Stock News Bulletin created successfully ({news_count} items included, {skipped_count} duplicates skipped).")
        return True

    except Exception as e:
        print(f"An error occurred: {e}")
        return False
    finally:
        if driver:
            try:
                driver.quit()
                print("Browser closed successfully")
            except:
                print("Warning: Could not close browser cleanly")

# --- Main Execution Block ---
if __name__ == "__main__":
    success = create_stocks_bulletin()
    if success:
        print(f"\n'{'Key_Stocks_to_Watch.docx'}' has been created in the current directory.")
    else:
        print("\nFailed to create the Stock News Bulletin.")
