import requests
import time
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt

# --- Functions for Word Document Formatting ---
# These are needed for the title's background shading.
try:
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    import docx.shared

    def add_shading_to_paragraph(paragraph, color="000000"):
        """Applies a background color shading to an entire paragraph."""
        shading_xml = f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color}"/>'
        pPr = paragraph._p.get_or_add_pPr()
        pPr.append(parse_xml(shading_xml))

except ImportError:
    print("Warning: python-docx Oxml features not available. Title background may not work.")
    def add_shading_to_paragraph(paragraph, color): pass


def create_filtered_market_bulletin(output_filename="Market_Bulletin_Filtered.docx"):
    """
    Scrapes market news, filters out items containing 'Moneycontrol',
    formats the top 10 remaining items, and saves them to a .docx file.

    Args:
        output_filename (str): The name of the output .docx file.

    Returns:
        bool: True if successful, False otherwise.
    """
    print("Fetching latest market news from MoneyControl...")
    
    try:
        url = "https://www.moneycontrol.com/news/business/markets/"
        session = requests.Session()
        time.sleep(2)  # Add a small delay
        headers = {
            'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/16.5 Safari/605.1.15',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.5',
            'Accept-Encoding': 'gzip, deflate',
            'Connection': 'keep-alive',
            'Upgrade-Insecure-Requests': '1',
            'Sec-Fetch-Dest': 'document',
            'Sec-Fetch-Mode': 'navigate',
            'Sec-Fetch-Site': 'none',
            'Sec-Fetch-User': '?1',
            'Cache-Control': 'max-age=0'
        }
        # First get the homepage to set cookies
        session.get("https://www.moneycontrol.com", headers=headers, timeout=10)
        time.sleep(1)  # Small delay between requests
        response = session.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        response.encoding = 'utf-8'

        content = response.content

        # Print response headers for debugging
        #print("Response Headers:", dict(response.headers))
        
        # Print response content length
        #print("Content Length:", len(response.content))

        soup = BeautifulSoup(content, 'html.parser')
        # First try 'cagetory'
        news_list = soup.find('ul', id='cagetory')
        
        # If not found, try other common containers
        if not news_list:
            print("Trying alternative selectors...")
            news_list = soup.find('ul', class_='article_listing')
            
        if not news_list:
            news_list = soup.find('div', class_='article-list')
            
        if not news_list:
            print("Could not find the news list on the page.")
            print("Page content preview:", soup.get_text()[:500])  # Print first 500 chars for debugging
            return False
            
        # Fetch more headlines than we need, to account for filtering
        headlines_items = news_list.find_all('li', class_='clearfix', limit=20)
        if not headlines_items:
            print("Could not find any news items."); return False

        print(f"Filtering news and creating '{output_filename}'...")
        document = Document()
        
        # --- Add Title ---
        title_paragraph = document.add_paragraph()
        title_run = title_paragraph.add_run(" Market Bulletin ")
        font = title_run.font
        font.name = 'Arial Black'; font.size = Pt(20)
        if docx.shared: font.color.rgb = docx.shared.RGBColor(0xff, 0xff, 0xff)
        add_shading_to_paragraph(title_paragraph, color="000000")
        document.add_paragraph()

        # --- Filter and Add News Items ---
        news_count = 0
        for item in headlines_items:
            # Stop after we have found 10 good headlines
            if news_count >= 10:
                break

            headline_tag = item.find('h2')
            summary_tag = item.find('p')
            
            if headline_tag and headline_tag.a and summary_tag:
                headline_text = headline_tag.a.get_text(strip=True)
                summary_text = summary_tag.get_text(strip=True)
                
                # --- THE FILTERING LOGIC IS HERE ---
                # Check if 'moneycontrol' (case-insensitive) is in either text
                if 'moneycontrol' in headline_text.lower() or 'moneycontrol' in summary_text.lower():
                    print(f"  > Skipping self-promotional article: '{headline_text[:50]}...'")
                    continue # Skip to the next item in the loop
                
                # --- If the filter passes, add the content to the document ---
                
                # Add numbered headline in bold
                # Add numbered headline with custom formatting
                p_headline = document.add_paragraph()
                number_run = p_headline.add_run(f"{news_count + 1}. ")  # Manual numbering
                number_run.bold = True
                number_run.font.size = Pt(12)
                
                p_headline.paragraph_format.left_indent = Inches(0.25)
                p_headline.paragraph_format.first_line_indent = Inches(-0.25)
                runner = p_headline.add_run(f"{headline_text}")
                runner.bold = True
                runner.font.size = Pt(12)
                
                # Add styled summary text below
                p_summary = document.add_paragraph()
                p_summary.paragraph_format.left_indent = Inches(0.5)
                summary_run = p_summary.add_run(summary_text)
                summary_run.italic = True  # Make summary italic
                summary_run.font.size = Pt(11)  # Slightly smaller font
                summary_run.font.color.rgb = docx.shared.RGBColor(89, 89, 89)  # Gray color
                
                # Adjust spacing
                p_headline.paragraph_format.space_before = Pt(12)
                p_headline.paragraph_format.space_after = Pt(6)
                p_summary.paragraph_format.space_before = Pt(0)
                p_summary.paragraph_format.space_after = Pt(18)  # More space between items
                
                # Increment the counter of news items we've added
                news_count += 1
        
        if news_count == 0:
            print("Could not find any suitable news after filtering.")
            return False

        document.save(output_filename)
        print(f"Filtered Market Bulletin with {news_count} items created successfully.")
        #return True
        return "\n".join([para.text for para in document.paragraphs if para.text.strip()])

    except Exception as e:
        print(f"An error occurred: {e}")
        return False

# --- Main Execution Block ---
if __name__ == "__main__":
    success = create_filtered_market_bulletin()
    if success:
        print(f"\n'{'Market_Bulletin_Filtered.docx'}' has been saved in the current directory.")
    else:
        print("\nFailed to create the Market Bulletin.")
