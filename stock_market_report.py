import os
import time
import io
import yfinance as yf
import pandas as pd
import requests
import re
from bs4 import BeautifulSoup
from PIL import Image
import undetected_chromedriver as uc
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service as ChromeService
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from datetime import datetime
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor, black, grey, white, lightgrey
from docx import Document
from docx.shared import Inches, Pt, RGBColor
import mplfinance as mpf
from reportlab.platypus import Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from nsepython import nse_optionchain_scrapper
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def add_shading_to_paragraph(paragraph, color="000000"):
    """Applies a background color shading to an entire paragraph."""
    shading_xml = f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color}"/>'
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(parse_xml(shading_xml))

# Create output directory
os.makedirs("CodeOutput", exist_ok=True)

def get_key_stocks_to_watch():
    """Generate Key Stocks to Watch Report"""
    print("Fetching latest stocks news from Groww...")
    driver = None
    wait = None
    
    try:
        # Initialize browser
        print("Initializing browser...")
        options = uc.ChromeOptions()
        
        # Core settings for stability
        options.add_argument('--headless')
        options.add_argument('--no-sandbox')
        options.add_argument('--disable-dev-shm-usage')
        options.add_argument('--disable-gpu')
        options.add_argument('--window-size=1920,1080')
        options.add_argument('--disable-web-security')
        options.add_argument('--ignore-certificate-errors')
        options.add_argument('--disable-extensions')
        options.add_argument('--disable-notifications')
        options.add_argument('--enable-features=NetworkService,NetworkServiceInProcess')
        
        # Initialize driver with retry mechanism
        max_retries = 3
        retry_count = 0
        retry_delay = 5  # seconds between retries
        
        while retry_count < max_retries:
            try:
                if retry_count > 0:
                    print(f"Retry attempt {retry_count} of {max_retries}...")
                    time.sleep(retry_delay)
                    retry_delay *= 2
                    
                if driver:
                    try:
                        driver.close()
                        driver.quit()
                    except:
                        pass
                    driver = None
                
                # Create fresh ChromeOptions for each attempt
                options = uc.ChromeOptions()
                options.add_argument('--headless')
                options.add_argument('--no-sandbox')
                options.add_argument('--disable-dev-shm-usage')
                options.add_argument('--disable-gpu')
                options.add_argument('--window-size=1920,1080')
                options.add_argument('--disable-web-security')
                options.add_argument('--ignore-certificate-errors')
                options.add_argument('--disable-extensions')
                options.add_argument('--disable-notifications')
                options.add_argument('--enable-features=NetworkService,NetworkServiceInProcess')
                
                driver = uc.Chrome(version_main=137, options=options, use_subprocess=True)
                driver.set_page_load_timeout(30)
                driver.set_script_timeout(30)
                wait = WebDriverWait(driver, 30)
                
                # Verify browser is working
                driver.get('about:blank')
                if not driver.current_url:
                    raise Exception("Browser failed to initialize properly")
                    
                print("Browser initialized successfully")
                break
                
            except Exception as e:
                retry_count += 1
                print(f"Browser initialization attempt {retry_count} failed: {str(e)}")
                
                if retry_count == max_retries:
                    raise Exception(f"Failed to initialize browser after {max_retries} retries")
                    
        # Load page and wait for items with validation
        url = "https://groww.in/market-news/stocks"
        print("Loading page...")
        try:
            driver.get(url)
            wait.until(EC.presence_of_element_located((By.CLASS_NAME, 'smnli671ItemContainer')))
            print("Page loaded successfully")
        except Exception as e:
            print(f"Error loading page: {e}")
            if driver:
                driver.quit()
            return False
        
        # Scroll to load all items with validation
        print("Loading more news items...")
        SCROLL_PAUSE_TIME = 2
        scroll_count = 0
        max_scrolls = 2
        
        while scroll_count < max_scrolls:
            try:
                # Verify browser is still responsive
                driver.current_url
                
                # Scroll down
                driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
                time.sleep(SCROLL_PAUSE_TIME)
                scroll_count += 1
                print(f"Scroll {scroll_count}/{max_scrolls}")
                
                # Try to click "Load More" button if it exists
                try:
                    load_more = wait.until(EC.presence_of_element_located((By.CLASS_NAME, "sm403InfiniteLoaderContainer")))
                    if load_more and load_more.is_displayed():
                        driver.execute_script("arguments[0].scrollIntoView(true);", load_more)
                        time.sleep(1)
                except:
                    pass  # No load more button found
                    
            except Exception as e:
                print(f"Error during scrolling: {e}")
                break
        
        # Get all loaded items
        try:
            soup = BeautifulSoup(driver.page_source, 'html.parser')
            news_items = soup.find_all('div', {'class': lambda x: x and 'ItemContainer' in x})
            
            if not news_items:
                print("Could not find news items.")
                return False
                
        except Exception as e:
            print(f"Error parsing page content: {e}")
            return False
            
        print(f"Found {len(news_items)} news items to process")

        print(f"Filtering news and creating 'Key_Stocks_to_Watch.docx'...")
        # Create document
        output_file = os.path.join("CodeOutput", "Key_Stocks_to_Watch.docx")
        try:
            document = Document()
            
            # Check if file is writable
            with open(output_file, 'a') as test_write:
                pass
                
        except Exception as e:
            print(f"Error creating document: {e}")
            if driver:
                driver.quit()
            return False
        
        # Add title
        title_paragraph = document.add_paragraph()
        title_run = title_paragraph.add_run(" Key Stocks to Watch ")
        font = title_run.font
        font.name = 'Arial Black'
        font.size = Pt(20)
        font.color.rgb = RGBColor(255, 255, 255)  # White text color
        add_shading_to_paragraph(title_paragraph, color="000000")
        document.add_paragraph()

        # Process news items
        seen_stocks = set()  # Track unique stock names
        news_count = 0  # Track number of unique news added
        skipped_count = 0  # Track duplicates
        
        for idx, item in enumerate(news_items, 1):
            print(f"\rProcessing news {idx}/{len(news_items)}...", end="")
            if news_count >= 15:  # Get 15 unique items
                break
            try:
                # Extract news details
                # Get the header container
                header_div = item.find('div', class_='smnli671BoxHeaderText')
                if not header_div:
                    header_div = item.find('div', {'class': lambda x: x and 'BoxHeaderText' in x})
                
                if header_div:
                    source = header_div.find('div').text.strip()
                    time_element = header_div.find('time')
                    timestamp = time_element.text.strip() if time_element else ""
                else:
                    print("Warning: Could not find header for news item", news_count + 1)
                    source = "Unknown Source"
                    timestamp = ""
                
                # Get the headline
                headline_div = item.find('div', class_='smnli671BoxItemTitle')
                if not headline_div:
                    headline_div = item.find('div', {'class': lambda x: x and 'BoxItemTitle' in x})
                
                if headline_div:
                    headline = headline_div.text.strip()
                else:
                    print("Warning: Could not find headline for news item", news_count + 1)
                    continue
                
                # Find stock details with flexible class matching
                stock_container = item.find('span', class_='smnli671MarketNewsCompName') or \
                                item.find('span', {'class': lambda x: x and 'MarketNewsCompName' in x})
                if stock_container:
                    stock_name = stock_container.text.strip()
                    if stock_name in seen_stocks:
                        skipped_count += 1
                        print(f"  > Skipping duplicate stock: {stock_name}")
                        continue
                    seen_stocks.add(stock_name)
                    
                    # Get price change by looking at all sibling elements
                    price_change = "N/A"
                    parent = stock_container.parent
                    if parent:
                        for span in parent.find_all('span'):
                            if 'content' in str(span.get('class', [])):
                                price_text = span.text.strip()
                                classes = span.get('class', [])
                                
                                if any('Positive' in c for c in classes):
                                    if not price_text.startswith('+'):
                                        price_change = f"+{price_text}"
                                    else:
                                        price_change = price_text
                                    break
                                elif any('Negative' in c for c in classes):
                                    if not price_text.startswith('-'):
                                        price_change = f"-{price_text}"
                                    else:
                                        price_change = price_text
                                    break
                else:
                    stock_name = ""
                    price_change = ""
                
                news_count += 1  # Increment counter for unique news
                
                # Add numbered headline
                p_headline = document.add_paragraph()
                p_headline.paragraph_format.left_indent = Inches(0.25)
                p_headline.paragraph_format.first_line_indent = Inches(-0.25)
                p_headline.paragraph_format.space_before = Pt(12)
                p_headline.paragraph_format.space_after = Pt(6)

                number_run = p_headline.add_run(f"{news_count}. ")
                number_run.bold = True
                number_run.font.size = Pt(12)
                
                headline_run = p_headline.add_run(f"{headline}")
                headline_run.bold = True
                headline_run.font.size = Pt(12)
                
                # Add source and time with proper spacing
                p_source = document.add_paragraph()
                p_source.paragraph_format.left_indent = Inches(0.5)
                p_source.paragraph_format.space_before = Pt(3)
                p_source.paragraph_format.space_after = Pt(3)
                source_run = p_source.add_run(f"{source} • {timestamp}")
                source_run.italic = True
                source_run.font.size = Pt(10)
                source_run.font.color.rgb = RGBColor(89, 89, 89)
                
                # Add stock details if available
                if stock_name:
                    p_stock = document.add_paragraph()
                    p_stock.paragraph_format.left_indent = Inches(0.5)
                    p_stock.paragraph_format.space_before = Pt(3)
                    p_stock.paragraph_format.space_after = Pt(12)
                    stock_run = p_stock.add_run(f"{stock_name} ({price_change})")
                    stock_run.bold = True
                    stock_run.font.size = Pt(11)
                    if price_change.startswith('+'):
                        stock_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
                    elif price_change.startswith('-'):
                        stock_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
                else:
                    # Add extra space after source if no stock details
                    p_source.paragraph_format.space_after = Pt(12)
                
            except Exception as e:
                print(f"Error processing news item {news_count + 1}: {e}")
                continue

        try:
            document.save(output_file)
            print(f"Stock News Bulletin created successfully ({news_count} items included, {skipped_count} duplicates skipped).")
        except Exception as e:
            print(f"Error saving document: {e}")
            return False
        try:
            if driver:
                driver.close()
                driver.quit()
                print("Browser closed successfully")
                print("Function 8 - Key Stocks - Successful")
        except Exception as e:
            print(f"Warning: Could not close browser cleanly - {e}")
        return True

    except Exception as e:
        print(f"An error occurred: {e}")
        print("Function 8 - Key Stocks - Not Successful")
        if 'driver' in locals():
            driver.quit()
        return False

def get_nifty_summary():
    """Generate NIFTY50 Summary Report"""
    try:
        # Original print statements commented out
        # print("Fetching enhanced Nifty 50 data...")
        
        ticker = yf.Ticker("^NSEI")
        hist_1y = ticker.history(period="1y")
        hist_2d = ticker.history(period="2d")

        if hist_1y.empty or hist_2d.empty:
            # print("Could not download data. Check ticker or internet connection.")
            return False

        latest_day = hist_2d.iloc[-1]
        prev_day = hist_2d.iloc[-2]

        data = {
            "current_price": latest_day['Close'],
            "prev_close": prev_day['Close'],
            "open": latest_day['Open'],
            "intraday_high": latest_day['High'],
            "intraday_low": latest_day['Low'],
            "volume_lakhs": latest_day['Volume'] / 100000,
            "fifty_two_week_high": hist_1y['High'].max(),
            "fifty_two_week_low": hist_1y['Low'].min(),
        }
        data['change'] = data['current_price'] - data['prev_close']
        data['change_percent'] = (data['change'] / data['prev_close']) * 100

        # Generate PDF report
        pdf_file = os.path.join("CodeOutput", "Market_Report_Dashboard_Nifty50.pdf")
        c = canvas.Canvas(pdf_file, pagesize=(8.5*inch, 5*inch))
        width, height = (8.5*inch, 5*inch)

        color_indigo = HexColor("#4f46e5")
        color_green = HexColor("#16a34a")
        color_red = HexColor("#dc2626")

        c.setFont("Helvetica-Bold", 18)
        c.setFillColor(black)
        c.drawString(0.5 * inch, height - 0.7 * inch, "NIFTY 50")

        c.setFont("Helvetica-Bold", 42)
        c.drawString(0.5 * inch, height - 1.3 * inch, f"{data['current_price']:,.2f}")

        x_change = 3.5 * inch
        y_change = height - 1.2 * inch
        
        if data['change'] >= 0:
            c.setFillColor(color_green)
            change_text = f"▲ {data['change']:.2f} ({data['change_percent']:.2f}%)"
        else:
            c.setFillColor(color_red)
            change_text = f"▼ {abs(data['change']):.2f} ({abs(data['change_percent']):.2f}%)"
        
        c.setFont("Helvetica-Bold", 16)
        c.drawString(x_change, y_change, change_text)
        
        y_grid = height - 2.1 * inch
        
        # Helper function for drawing data points
        def draw_data_point(c, x, y, label, value, value_color):
            c.setFont("Helvetica", 9)
            c.setFillColor(grey)
            c.drawString(x, y, label)
            c.setFont("Helvetica-Bold", 14)
            c.setFillColor(value_color)
            c.drawString(x, y - 18, value)

        draw_data_point(c, 0.5 * inch, y_grid, "Prev. Close", f"{data['prev_close']:,.2f}", color_indigo)
        draw_data_point(c, 2.2 * inch, y_grid, "Open", f"{data['open']:,.2f}", color_indigo)
        draw_data_point(c, 3.9 * inch, y_grid, "Volume (Lakhs)", f"{data['volume_lakhs']:,.2f}", color_indigo)

        def draw_slider_refined(c, x, y, width, height, label, low_val, high_val, current_val):
            c.saveState()
            c.setFont("Helvetica-Bold", 14)
            c.setFillColor(black)
            c.drawString(x, y, label)

            y_bar = y - 35
            bar_radius = height / 2

            path = c.beginPath()
            path.roundRect(x, y_bar - bar_radius, width, height, bar_radius)
            c.clipPath(path, stroke=0, fill=0)
            c.linearGradient(x, y_bar, x + width, y_bar, (HexColor("#d90429"), HexColor("#f8b24f"), HexColor("#8ac926")), extend=False)

            c.restoreState()

            range_val = high_val - low_val
            position_ratio = (current_val - low_val) / range_val if range_val != 0 else 0.5
            position_ratio = max(0, min(1, position_ratio))
            marker_x = x + width * position_ratio

            label_text = f"{current_val:,.2f}"
            label_width = c.stringWidth(label_text, "Helvetica", 9) + 10
            label_height = 14
            label_x = marker_x - (label_width / 2)
            label_y = y_bar + 10

            c.setFillColor(white)
            c.setStrokeColor(lightgrey)
            c.roundRect(label_x, label_y, label_width, label_height, label_height/2, stroke=1, fill=1)
            
            c.setFillColor(black)
            c.setFont("Helvetica", 9)
            c.drawCentredString(marker_x, label_y + 4, label_text)

            c.setFillColor(white)
            c.setStrokeColor(grey)
            c.setLineWidth(1)
            c.circle(marker_x, y_bar, 6, stroke=1, fill=1)

            y_lowhigh = y_bar - 25
            c.setFont("Helvetica", 9)
            c.setFillColor(grey)
            c.drawString(x, y_lowhigh, "Low")
            c.drawRightString(x + width, y_lowhigh, "High")

            c.setFont("Helvetica", 12)
            c.setFillColor(black)
            c.drawString(x, y_lowhigh - 15, f"{low_val:,.2f}")
            c.drawRightString(x + width, y_lowhigh - 15, f"{high_val:,.2f}")

        slider_y = height - 3.2 * inch
        slider_width = 3.5 * inch
        slider_height = 10

        draw_slider_refined(c, 0.5 * inch, slider_y, slider_width, slider_height, "52 Week",
                    data['fifty_two_week_low'], data['fifty_two_week_high'], data['current_price'])
        
        draw_slider_refined(c, 4.5 * inch, slider_y, slider_width, slider_height, "Intraday",
                    data['intraday_low'], data['intraday_high'], data['current_price'])

        c.save()
        
        print("Function 1 - NIFTY50 Summary - Successful")
        return True

    except Exception as e:
        print("Function 1 - NIFTY50 Summary - Not Successful")
        return False

def get_sgx_nifty():
    """Generate SGX Nifty Analysis"""
    try:
        options = webdriver.ChromeOptions()
        options.add_argument("--headless")
        options.add_argument("--window-size=1280,1000")
        options.add_argument("--disable-gpu")
        options.add_argument("--no-sandbox")
        
        service = ChromeService(ChromeDriverManager().install())
        driver = webdriver.Chrome(service=service, options=options)
        
        url = 'https://sgxnifty.org/'
        driver.get(url)
        
        wait = WebDriverWait(driver, 30)
        
        # Find the main container
        container_selector = ".grid.col-940.align-center"
        container_element = wait.until(
            EC.visibility_of_element_located((By.CSS_SELECTOR, container_selector))
        )
        time.sleep(2)  # Allow dynamic values to load

        # Take screenshot
        screenshot_data = container_element.screenshot_as_png

        # Find tables for cropping
        tables = container_element.find_elements(By.CLASS_NAME, "main-table-div")
        if len(tables) < 2:
            return False
            
        last_element_to_keep = tables[1]
        crop_height = last_element_to_keep.location['y'] + last_element_to_keep.size['height'] - 130
        
        # Process and save image
        img = Image.open(io.BytesIO(screenshot_data))
        crop_box = (0, 0, img.width, crop_height)
        cropped_img = img.crop(crop_box)

        output_file = os.path.join("CodeOutput", "sgx_nifty.png")
        cropped_img.save(output_file, "PNG")

        driver.quit()
        
        print("Function 11 - SGX Nifty Analysis - Successful")
        return True

    except Exception as e:
        print("Function 11 - SGX Nifty Analysis - Not Successful")
        if 'driver' in locals():
            driver.quit()
        return False

def get_silver_rates():
    """Generate Silver Rates Analysis"""
    try:
        url = 'https://www.goodreturns.in/silver-rates/chennai.html'
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }

        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()

        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Get Today's Silver Prices
        price_container = soup.find('div', class_='gold-rate-container')
        if not price_container:
            return False
        
        todays_rates_divs = price_container.find_all('div', class_='gold-each-container')
        if len(todays_rates_divs) < 2:
            return False

        def parse_price_box(box):
            bottom_div = box.find('div', class_='gold-bottom')
            p_tags = bottom_div.find_all('p')
            price_str = p_tags[0].text if len(p_tags) > 0 else "0"
            change_str = p_tags[1].text if len(p_tags) > 1 else "0"
            price = float(re.sub(r'[₹,]', '', price_str))
            change_val = float(re.sub(r'[₹,+-]', '', change_str))
            if '-' in change_str:
                change_val = -change_val
            return {"price": price, "change": change_val}

        today_per_gram = parse_price_box(todays_rates_divs[0])
        today_per_kg = parse_price_box(todays_rates_divs[1])

        # Get Historical Data
        all_tables = soup.find_all('table')
        history_table = None
        
        for table in all_tables:
            headers = [th.get_text(strip=True) for th in table.find_all('th')]
            if headers == ['Date', '10 gram', '100 gram', '1 Kg']:
                history_table = table
                break
        
        if not history_table:
            return False
            
        history_rows = history_table.find('tbody').find_all('tr')
        
        historical_data = []
        for row in history_rows:
            cols = row.find_all('td')
            if len(cols) == 4:
                date = cols[0].get_text(strip=True)
                price_10g = cols[1].get_text(strip=True)
                price_100g = cols[2].get_text(strip=True)
                price_1kg_full = ' '.join(cols[3].text.split())

                historical_data.append({
                    "date": date,
                    "price_10g": price_10g,
                    "price_100g": price_100g,
                    "price_1kg": price_1kg_full
                })

        # Save to file
        output_file = os.path.join("CodeOutput", "silver_rates.txt")
        with open(output_file, 'w') as f:
            f.write("CHENNAI SILVER RATE REPORT\n")
            f.write("=" * 60 + "\n\n")
            
            f.write("[Today's Silver Rate]\n")
            f.write("-" * 60 + "\n")
            
            price_g = today_per_gram['price']
            change_g = today_per_gram['change']
            arrow_g = "▲" if change_g >= 0 else "▼"
            f.write(f"Per Gram: ₹{price_g:,.2f} (Change: ₹{abs(change_g):.2f} {arrow_g})\n")
            
            price_kg = today_per_kg['price']
            change_kg = today_per_kg['change']
            arrow_kg = "▲" if change_kg >= 0 else "▼"
            f.write(f"Per Kg:   ₹{price_kg:,.0f} (Change: ₹{abs(change_kg):,.0f} {arrow_kg})\n\n")
            
            f.write("[Silver Rate in Chennai for Last 10 Days]\n")
            header = f"{'Date':<15} | {'10 gram':<15} | {'100 gram':<15} | {'1 Kg'}"
            f.write("-" * len(header) + "\n")
            f.write(header + "\n")
            f.write("-" * len(header) + "\n")
            
            for record in historical_data:
                f.write(f"{record['date']:<15} | {record['price_10g']:<15} | {record['price_100g']:<15} | {record['price_1kg']}\n")
            
            f.write("-" * len(header) + "\n")

        print("Function 14 - Silver Rates Analysis - Successful")
        return True

    except Exception as e:
        print("Function 14 - Silver Rates Analysis - Not Successful")
        return False

def get_fii_dii_data():
    """Generate FII/DII Analysis Report"""
    try:
        def arrow(val: float) -> str:
            return f"{val:,.2f} {'🟢↑' if val >= 0 else '🔴↓'}"

        # Fetch data from website
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
            "Accept-Language": "en-US,en;q=0.5",
            "Connection": "keep-alive",
        }

        all_data = []
        max_pages = 4
        page = 1

        while page <= max_pages:
            url = "https://groww.in/fii-dii-data" if page == 1 else f"https://groww.in/fii-dii-data?page={page}"
            resp = requests.get(url, headers=headers, timeout=10)
            resp.raise_for_status()
            
            soup = BeautifulSoup(resp.text, "html.parser")
            table = soup.find("table")
            if not table:
                break
            
            rows = table.select("tbody tr")
            for row in rows:
                cols = [td.get_text(strip=True).replace(",", "") for td in row.find_all("td")]
                if len(cols) >= 7:
                    date_str = cols[0]
                    fii_net = float(cols[3].replace("+", "").replace("−", "-").replace("–", "-"))
                    dii_net = float(cols[6].replace("+", "").replace("−", "-").replace("–", "-"))
                    parsed_date = datetime.strptime(date_str, "%d %b %Y")
                    all_data.append((parsed_date, fii_net, dii_net))
            page += 1

        if not all_data:
            return False

        df = pd.DataFrame(all_data, columns=["Date", "FII_Net", "DII_Net"])
        df = df.drop_duplicates(subset=['Date'], keep='first')
        df = df.sort_values("Date", ascending=False).reset_index(drop=True)

        results = []

        # Last 3 individual days
        for i in range(3):
            row = df.iloc[i]
            date_label = row['Date'].strftime("%d-%m-%Y")
            results.append({
                "Date": date_label,
                "FII": arrow(row['FII_Net']),
                "DII": arrow(row['DII_Net'])
            })

        # Last 7 days cumulative
        last_7 = df.head(7)
        results.append({
            "Date": "Last 7 Days",
            "FII": arrow(last_7['FII_Net'].sum()),
            "DII": arrow(last_7['DII_Net'].sum())
        })

        # Last 10 days cumulative
        last_10 = df.head(10)
        results.append({
            "Date": "Last 10 Days",
            "FII": arrow(last_10['FII_Net'].sum()),
            "DII": arrow(last_10['DII_Net'].sum())
        })

        # Save to file
        output_file = os.path.join("CodeOutput", "fii_dii_data.txt")
        with open(output_file, 'w') as f:
            f.write("FII/DII ACTIVITY SUMMARY\n")
            f.write("=" * 80 + "\n\n")
            
            for row in results:
                f.write(f"{row['Date']:<12}: ")
                f.write(f"FII = {row['FII']:<30} ")
                f.write(f"DII = {row['DII']}\n")
            
            f.write("\nNote: Values in Crores (₹)")

        print("Function 10 - FII/DII Analysis - Successful")
        return True

    except Exception as e:
        print("Function 10 - FII/DII Analysis - Not Successful")
        return False

def get_gold_rates():
    """Generate Gold Rates Analysis"""
    try:
        url = 'https://www.goodreturns.in/gold-rates/chennai.html'
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }

        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()

        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Get Today's Gold Prices
        price_container = soup.find('div', class_='gold-rate-container')
        if not price_container:
            return False
        
        todays_rates_divs = price_container.find_all('div', class_='gold-each-container')
        if len(todays_rates_divs) < 2:
            return False

        def parse_price_box(box):
            bottom_div = box.find('div', class_='gold-bottom')
            p_tags = bottom_div.find_all('p')
            price_str = p_tags[0].text if len(p_tags) > 0 else "0"
            change_str = p_tags[1].text if len(p_tags) > 1 else "0"
            price = float(re.sub(r'[₹,]', '', price_str))
            try:
                change_str = re.sub(r'[₹,]', '', change_str).strip()
                change_str = ''.join(change_str.split())
                if change_str.startswith('+'):
                    change_str = change_str[1:]
                elif change_str.startswith('- ') or change_str.startswith('-'):
                    change_str = f"-{change_str.replace('-', '').strip()}"
                change = float(change_str)
            except ValueError:
                change = 0.0
            return {"price": price, "change": change}

        today_24k = parse_price_box(todays_rates_divs[0])
        today_22k = parse_price_box(todays_rates_divs[1])

        # Get Historical Data
        tables = soup.find_all('table')
        history_table = None
        for table in tables:
            headers = [th.get_text(strip=True) for th in table.find_all('th')]
            if 'Date' in headers and '24K' in headers and '22K' in headers:
                history_table = table
                break
        
        if not history_table:
            return False
            
        history_rows = history_table.find_all('tr')
        if history_rows[0].find('th'):
            history_rows = history_rows[1:]
        
        historical_data = []
        for row in history_rows:
            cols = row.find_all('td')
            if len(cols) == 3:
                date = cols[0].text.strip()
                
                # 24K data
                td_24k = cols[1]
                price_24k = td_24k.get_text(strip=True).split('(')[0].strip()
                change_span_24k = td_24k.find('span')
                change_24k = ""
                change_color_24k = ""
                if change_span_24k:
                    change_text = change_span_24k.text.strip()
                    change_24k = change_text.strip('()')
                    change_24k = change_24k.replace('+', '')
                    if 'green-span' in change_span_24k.get('class', []):
                        change_color_24k = 'green'
                    elif 'red-span' in change_span_24k.get('class', []):
                        change_color_24k = 'red'
                
                # 22K data
                td_22k = cols[2]
                price_22k = td_22k.get_text(strip=True).split('(')[0].strip()
                change_span_22k = td_22k.find('span')
                change_22k = ""
                change_color_22k = ""
                if change_span_22k:
                    change_text = change_span_22k.text.strip()
                    change_22k = change_text.strip('()')
                    change_22k = change_22k.replace('+', '')
                    if 'green-span' in change_span_22k.get('class', []):
                        change_color_22k = 'green'
                    elif 'red-span' in change_span_22k.get('class', []):
                        change_color_22k = 'red'
                
                if date and price_24k and price_22k:
                    historical_data.append({
                        "date": date,
                        "price_24k": price_24k,
                        "change_24k": change_24k,
                        "change_color_24k": change_color_24k,
                        "price_22k": price_22k,
                        "change_22k": change_22k,
                        "change_color_22k": change_color_22k
                    })

        # Save to file
        output_file = os.path.join("CodeOutput", "gold_rates.txt")
        with open(output_file, 'w') as f:
            f.write("CHENNAI GOLD RATE REPORT\n")
            f.write("=" * 60 + "\n\n")
            
            f.write("[Today's Gold Rate (per gram)]\n")
            f.write("-" * 60 + "\n")
            price_24k = today_24k['price']
            change_24k = today_24k['change']
            arrow_24k = "▲" if change_24k >= 0 else "▼"
            f.write(f"24K Gold: ₹{price_24k:,.0f} (Change: ₹{abs(change_24k):.0f} {arrow_24k})\n")
            
            price_22k = today_22k['price']
            change_22k = today_22k['change']
            arrow_22k = "▲" if change_22k >= 0 else "▼"
            f.write(f"22K Gold: ₹{price_22k:,.0f} (Change: ₹{abs(change_22k):.0f} {arrow_22k})\n\n")
            
            f.write("[Gold Rate in Chennai for Last 10 Days (1 gram)]\n")
            f.write("-" * 104 + "\n")
            f.write(f"{'Date':<16} | {'24K Price':<42} | {'22K Price':<42}\n")
            f.write("-" * 104 + "\n")
            
            for record in historical_data:
                price_24k_text = f"{record['price_24k']}  ({record['change_24k']})"
                price_22k_text = f"{record['price_22k']}  ({record['change_22k']})"
                f.write(f"{record['date']:<16} | {price_24k_text:<42} | {price_22k_text:<42}\n")
            
            f.write("-" * 104 + "\n")

        print("Function 13 - Gold Rates Analysis - Successful")
        return True

    except Exception as e:
        print("Function 13 - Gold Rates Analysis - Not Successful")
        return False

def get_currency_rates():
    """Generate Currency Exchange Rates Analysis"""
    try:
        currencies = {
            "USD": {"name": "US Dollar", "country": "United States", "ticker": "USDINR=X"},
            "EUR": {"name": "Euro", "country": "Eurozone", "ticker": "EURINR=X"},
            "GBP": {"name": "British Pound", "country": "United Kingdom", "ticker": "GBPINR=X"},
            "SGD": {"name": "Singapore Dollar", "country": "Singapore", "ticker": "SGDINR=X"},
            "JPY": {"name": "Japanese Yen", "country": "Japan", "ticker": "JPYINR=X"},
            "AED": {"name": "UAE Dirham", "country": "U.A.E.", "ticker": "AEDINR=X"},
        }
        
        results = []
        
        for code, details in currencies.items():
            ticker = yf.Ticker(details['ticker'])
            info = ticker.info
            
            value = info.get('regularMarketPrice') or info.get('previousClose')
            
            if value is None:
                continue

            results.append({
                "Code": code,
                "Name": details['name'],
                "Country": details['country'],
                "Value": value
            })

        if not results:
            return False

        # Save to file
        output_file = os.path.join("CodeOutput", "currency_rates.txt")
        with open(output_file, 'w') as f:
            f.write("POPULAR CURRENCIES vs. INR\n")
            f.write("=" * 70 + "\n\n")
            f.write(f"{'Code':<5} | {'Country':<25} | {'Value (1 unit in INR)'}\n")
            f.write("-" * 70 + "\n")
            
            desired_order = ["USD", "JPY", "EUR", "SGD", "GBP", "AED"]
            data_map = {item['Code']: item for item in results}
            sorted_data = [data_map[code] for code in desired_order if code in data_map]
            
            for currency in sorted_data:
                code_str = currency['Code']
                country_str = currency['Country']
                value_str = f"₹{currency['Value']:.2f}"
                f.write(f"{code_str:<5} | {country_str:<25} | {value_str}\n")
            
            f.write("-" * 70)

        print("Function 15 - Currency Exchange Rates - Successful")
        return True

    except Exception as e:
        print("Function 15 - Currency Exchange Rates - Not Successful :(")
        return False

def get_global_markets():
    """Generate Global Markets Analysis"""
    try:
        indices = {
            "Dow Jones": "^DJI",
            "S&P 500": "^GSPC",
            "Nasdaq": "^IXIC",
            "FTSE 100": "^FTSE",
            "Hang Seng": "^HSI"
        }
        
        tickers_list = list(indices.values())
        
        # Fetch 5 days of data to ensure we have enough history
        data = yf.download(tickers=tickers_list, period="5d", auto_adjust=True)
        if data.empty:
            return False

        results = []
        ticker_to_name = {v: k for k, v in indices.items()}
        
        for ticker in tickers_list:
            ticker_data = data['Close'][ticker]
            ticker_data_cleaned = ticker_data.dropna()
            
            if len(ticker_data_cleaned) < 2:
                continue
            
            latest_close = ticker_data_cleaned.iloc[-1]
            prev_close = ticker_data_cleaned.iloc[-2]
            
            change = latest_close - prev_close
            change_percent = (change / prev_close) * 100
            
            results.append({
                "Name": ticker_to_name[ticker],
                "LTP": latest_close,
                "Change": change,
                "Change %": change_percent
            })

        if not results:
            return False

        # Save to file
        output_file = os.path.join("CodeOutput", "global_markets.txt")
        with open(output_file, 'w') as f:
            f.write("KEY GLOBAL INDICES\n")
            f.write("=" * 65 + "\n\n")
            f.write(f"{'Name':<15} | {'LTP':>15} | {'Change':>12} | {'Change %':>12}\n")
            f.write("-" * 65 + "\n")
            
            desired_order = ["Dow Jones", "Nasdaq", "S&P 500", "Hang Seng", "FTSE 100"]
            sorted_data = sorted(results, key=lambda x: desired_order.index(x['Name']))
            
            for index in sorted_data:
                ltp_str = f"{index['LTP']:,.2f}"
                change_str = f"{index['Change']:+.2f}"
                change_pct_str = f"{index['Change %']:+.2f}%"
                f.write(f"{index['Name']:<15} | {ltp_str:>15} | {change_str:>12} | {change_pct_str:>12}\n")
            
            f.write("-" * 65)

        print("Function 12 - Global Markets Analysis - Successful")
        return True

    except Exception as e:
        print("Function 12 - Global Markets Analysis - Not Successful")
        return False

def get_vix_analysis():
    """Generate India VIX Analysis"""
    print("Fetching VIX data and chart...")
    driver = None
    try:
        max_retries = 3
        retry_count = 0
        retry_delay = 5  # seconds between retries
        
        while retry_count < max_retries:
            try:
                if retry_count > 0:
                    print(f"Retry attempt {retry_count} of {max_retries}...")
                    time.sleep(retry_delay)
                    retry_delay *= 2
                
                if driver:
                    try:
                        driver.quit()
                    except:
                        pass
                    driver = None
                
                # Create fresh ChromeOptions for each attempt
                options = uc.ChromeOptions()
                options.add_argument("--headless")
                options.add_argument("--window-size=1920,1080")
                options.add_argument("--no-sandbox")
                options.add_argument('--disable-dev-shm-usage')
                options.add_argument('--disable-gpu')
                options.add_argument('--disable-extensions')
                options.add_argument('--disable-notifications')
                
                # Initialize driver with retry mechanism
                driver = uc.Chrome(version_main=137, options=options, use_subprocess=True)
                driver.set_page_load_timeout(30)
                driver.set_script_timeout(30)
                
                # Test browser is working
                driver.get('about:blank')
                if not driver.current_url:
                    raise Exception("Browser failed to initialize properly")
                    
                # Load VIX page
                page_url = 'https://groww.in/indices/india-vix'
                driver.get(page_url)
                wait = WebDriverWait(driver, 30)
                time.sleep(5)
                
                # Get VIX data
                table_rows = driver.find_elements(By.TAG_NAME, "tr")
                current_value = None
                prev_close = None
                open_value = None
                
                # Find VIX value
                for row in table_rows:
                    row_text = row.text
                    if 'INDIA VIX' in row_text:
                        parts = row_text.split('\n') if '\n' in row_text else row_text.split()
                        for i, part in enumerate(parts):
                            if part == 'INDIA VIX':
                                current_value = float(parts[i+1])
                                break
                
                # Find Prev.Close and Open values
                all_text = driver.find_element(By.TAG_NAME, "body").text
                
                prev_close_idx = all_text.find("Prev. Close")
                if prev_close_idx != -1:
                    lines = all_text[prev_close_idx:].split('\n')
                    for line in lines[:3]:
                        if line.replace('.', '').replace(',', '').strip().isdigit():
                            prev_close = float(line.strip())
                            break
                
                open_idx = all_text.find("Open")
                if open_idx != -1:
                    lines = all_text[open_idx:].split('\n')
                    for line in lines[:3]:
                        if line.replace('.', '').replace(',', '').strip().isdigit():
                            open_value = float(line.strip())
                            break
                
                if None in (current_value, prev_close, open_value):
                    raise Exception("Failed to get VIX values")
                
                print(f"VIX values found - Current: {current_value}, Prev Close: {prev_close}, Open: {open_value}")
                
                # Calculate changes
                change_value = current_value - open_value
                change_percentage = (change_value / open_value) * 100
                
                # Save VIX data to file
                output_file = os.path.join("CodeOutput", "india_vix.txt")
                with open(output_file, 'w') as f:
                    f.write("INDIA VIX ANALYSIS\n")
                    f.write("=" * 40 + "\n\n")
                    f.write(f"Current Value:    {current_value:.2f}\n")
                    f.write(f"Previous Close:   {prev_close:.2f}\n")
                    f.write(f"Open:            {open_value:.2f}\n")
                    f.write(f"Change:          {change_value:+.2f} ({change_percentage:+.2f}%)\n")
                
                # Capture chart
                driver.execute_script("""
                    document.querySelectorAll('.modal, .overlay').forEach(el => el.style.display = 'none');
                    document.documentElement.setAttribute('data-theme', 'light');
                """)
                
                timeframe_buttons = driver.find_elements(By.XPATH, "//button[text()='1M']")
                if timeframe_buttons:
                    driver.execute_script("arguments[0].click();", timeframe_buttons[0])
                    time.sleep(2)
                
                temp_screenshot = os.path.join("CodeOutput", "temp_vix.png")
                chart_file = os.path.join("CodeOutput", "india_vix_chart.png")
                
                driver.save_screenshot(temp_screenshot)
                
                img = Image.open(temp_screenshot)
                width, height = img.size
                
                left = width * 0.15
                right = width * 0.65
                top = height * 0.25
                bottom = height * 0.80
                
                chart_area = img.crop((left, top, right, bottom))
                chart_area.save(chart_file)
                
                os.remove(temp_screenshot)
                print("Function 9 - India VIX Analysis - Successful")
                return True
                
            except Exception as e:
                retry_count += 1
                print(f"Attempt {retry_count} failed: {str(e)}")
                if retry_count == max_retries:
                    raise
        
        return False
        
    except Exception as e:
        print(f"VIX Analysis failed: {str(e)}")
        return False
        
    finally:
        if driver:
            try:
                driver.quit()
                print("Browser closed successfully")
            except:
                print("Warning: Could not close browser cleanly")

def get_market_news():
    """Get Top 10 Market News"""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/16.5 Safari/605.1.15',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.5',
            'Accept-Encoding': 'gzip, deflate',
            'Connection': 'keep-alive'
        }
        
        url = "https://www.livemint.com/market/stock-market-news"
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()

        soup = BeautifulSoup(response.content, 'html.parser')
        headlines_items = []
        
        # Find all h2 headings first
        h2_items = soup.find_all('h2', class_='headline')
        if h2_items:
            headlines_items.extend(h2_items)
            
        # Also look for article headlines in list items
        article_items = soup.find_all('div', class_='listingNew')
        if article_items:
            headlines_items.extend(article_items)

        if not headlines_items:
            return False

        # Create document
        doc = Document()
        
        # Add title
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(" Market News Bulletin ")
        font = title_run.font
        font.name = 'Arial Black'
        font.size = Pt(20)
        font.color.rgb = RGBColor(255, 255, 255)  # White text color
        add_shading_to_paragraph(title_paragraph, color="000000")
        doc.add_paragraph()

        # Add news items
        news_count = 0
        for item in headlines_items[:10]:  # Only process first 10 items
            headline_text = item.get_text(strip=True)
            
            if headline_text:
                # Add numbered headline
                p_headline = doc.add_paragraph()
                number_run = p_headline.add_run(f"{news_count + 1}. ")
                number_run.bold = True
                number_run.font.size = Pt(12)
                
                p_headline.paragraph_format.left_indent = Inches(0.25)
                p_headline.paragraph_format.first_line_indent = Inches(-0.25)
                runner = p_headline.add_run(f"{headline_text}")
                runner.bold = True
                runner.font.size = Pt(12)
                
                p_headline.paragraph_format.space_before = Pt(12)
                p_headline.paragraph_format.space_after = Pt(12)
                
                news_count += 1

        # Save document
        output_file = os.path.join("CodeOutput", "Market_Bulletin.docx")
        doc.save(output_file)
        
        print("Function 7 - Top 10 Market News - Successful")
        return True

    except Exception as e:
        print("Function 7 - Top 10 Market News - Not Successful")
        return False

def get_nifty_oi():
    """Generate NIFTY50 Open Interest Analysis"""
    try:
        options = webdriver.ChromeOptions()
        options.add_argument("--headless")
        options.add_argument("--window-size=1920,1200")
        options.add_argument("--disable-gpu")
        options.add_argument("--no-sandbox")
        
        service = ChromeService(ChromeDriverManager().install())
        driver = webdriver.Chrome(service=service, options=options)
        
        url = 'https://upstox.com/fno-discovery/open-interest-analysis/nifty-oi/'
        driver.get(url)

        wait = WebDriverWait(driver, 30)
        
        # Wait for real data
        spot_price_xpath = "//div[p[text()='Spot']]/p[2]"
        wait.until(
            lambda driver: driver.find_element(By.XPATH, spot_price_xpath).text != '--/--'
        )

        # Extract data
        summary_container = driver.find_element(By.CSS_SELECTOR, "div.gap-8.py-4")
        data_sections = summary_container.find_elements(By.CSS_SELECTOR, "div.flex-col")
        
        extracted_data = {}
        for section in data_sections:
            p_tags = section.find_elements(By.TAG_NAME, 'p')
            if len(p_tags) == 2:
                label = p_tags[0].text
                value = p_tags[1].text
                extracted_data[label] = value

        spot_price_str = extracted_data.get('Spot', '0')
        total_calls_oi = extracted_data.get('Total Calls', '0 L')
        total_puts_oi = extracted_data.get('Total Puts', '0 L')
        
        spot_price = float(spot_price_str.replace(',', ''))
        
        # Save OI data to file
        output_file = os.path.join("CodeOutput", "nifty_oi.txt")
        with open(output_file, 'w') as f:
            f.write("NIFTY OPEN INTEREST ANALYSIS\n")
            f.write("=" * 50 + "\n\n")
            f.write(f"Spot Price:      {spot_price:,.2f}\n")
            f.write(f"Total Call OI:   {total_calls_oi}\n")
            f.write(f"Total Put OI:    {total_puts_oi}\n")
        
        # Capture chart
        chart_container_selector = "div.recharts-responsive-container"
        chart_element = wait.until(
            EC.visibility_of_element_located((By.CSS_SELECTOR, chart_container_selector))
        )
        
        time.sleep(2)
        
        temp_screenshot = os.path.join("CodeOutput", "temp_oi.png")
        chart_file = os.path.join("CodeOutput", "nifty_oi_chart.png")
        
        driver.save_screenshot(temp_screenshot)
        
        img = Image.open(temp_screenshot)
        width, height = img.size
        
        # Crop chart area
        left = width * 0.15
        right = width * 0.67
        top = height * 0.20
        bottom = height * 0.72
        
        chart_area = img.crop((left, top, right, bottom))
        chart_area.save(chart_file)
        
        os.remove(temp_screenshot)
        
        driver.quit()
        
        print("Function 6 - NIFTY50 Open Interest Analysis - Successful")
        return True

    except Exception as e:
        print("Function 6 - NIFTY50 Open Interest Analysis - Not Successful")
        if 'driver' in locals():
            driver.quit()
        return False

def get_nifty_pcr():
    """Generate NIFTY50 PCR Analysis"""
    try:
        # Get current PCR
        pcr_data = {}
        chain = nse_optionchain_scrapper('NIFTY')
        if not chain:
            return False

        total_pe_oi = sum(float(strike.get('PE', {}).get('openInterest', 0))
                         for strike in chain['records']['data'] if 'PE' in strike)
        total_ce_oi = sum(float(strike.get('CE', {}).get('openInterest', 0))
                         for strike in chain['records']['data'] if 'CE' in strike)

        if total_ce_oi == 0:
            return False

        current_pcr = total_pe_oi / total_ce_oi
        today_str = datetime.now().strftime('%d-%m-%Y')
        
        # Save PCR to file
        output_file = os.path.join("CodeOutput", "nifty_pcr.txt")
        with open(output_file, 'w') as f:
            f.write(f"NIFTY PCR Analysis\n")
            f.write("=" * 40 + "\n\n")
            f.write(f"Current PCR ({today_str}): {current_pcr:.2f}\n")
            f.write(f"Total Put OI: {total_pe_oi:,.0f}\n")
            f.write(f"Total Call OI: {total_ce_oi:,.0f}\n")

        # Capture PCR chart
        chart_url = "https://upstox.com/fno-discovery/open-interest-analysis/nifty-pcr/"
        chrome_options = Options()
        chrome_options.add_argument("--headless")
        chrome_options.add_argument("--window-size=1920,1080")
        chrome_options.add_argument("--no-sandbox")
        
        driver = webdriver.Chrome(options=chrome_options)
        
        try:
            driver.get(chart_url)
            wait = WebDriverWait(driver, 20)
            
            # Wait for chart to load
            wait.until(
                EC.presence_of_element_located((By.CLASS_NAME, "recharts-responsive-container"))
            )
            
            driver.set_window_size(1920, 1080)
            driver.execute_script("""
                document.querySelectorAll('.modal, .overlay').forEach(el => el.style.display = 'none');
                document.documentElement.setAttribute('data-theme', 'light');
            """)
            
            time.sleep(3)
            
            # Take and crop screenshot
            temp_screenshot = os.path.join("CodeOutput", "temp_pcr.png")
            pcr_chart = os.path.join("CodeOutput", "nifty_pcr_chart.png")
            
            driver.save_screenshot(temp_screenshot)
            
            img = Image.open(temp_screenshot)
            width, height = img.size
            
            # Crop chart area
            left = width * 0.15
            right = width * 0.67
            top = height * 0.30
            bottom = height * 0.81
            
            chart_area = img.crop((left, top, right, bottom))
            chart_area.save(pcr_chart)
            
            # Clean up
            os.remove(temp_screenshot)
            
        finally:
            driver.quit()

        print("Function 5 - NIFTY50 PCR - Successful")
        return True

    except Exception as e:
        print("Function 5 - NIFTY50 PCR - Not Successful")
        return False

def get_nifty_heatmap():
    """Generate NIFTY50 Heatmap"""
    try:
        # Initialize stealth browser
        options = uc.ChromeOptions()
        options.add_argument("--start-maximized")
        options.add_argument("--no-sandbox")
        options.add_argument("--disable-dev-shm-usage")
        options.add_argument("--disable-popup-blocking")
        
        # Ensure compatibility by using the correct ChromeDriver version
        driver = uc.Chrome(version_main=137, options=options, use_subprocess=True)
        wait = WebDriverWait(driver, 30)
        
        try:
            # Load page
            url = 'https://in.tradingview.com/heatmap/stock/#%7B%22dataSource%22%3A%22NIFTY50%22%2C%22blockColor%22%3A%22change%22%2C%22blockSize%22%3A%22market_cap_basic%22%2C%22grouping%22%3A%22no_group%22%7D'
            driver.get(url)
            time.sleep(5)
            
            # Click Fullscreen Button
            fullscreen_button = wait.until(
                EC.element_to_be_clickable((By.CSS_SELECTOR, "[data-qa-id='heatmap-top-bar_fullscreen']"))
            )
            fullscreen_button.click()
            time.sleep(3)

            # Click Display Settings Button
            settings_button = wait.until(
                EC.element_to_be_clickable((By.CSS_SELECTOR, "[data-qa-id='heatmap-top-bar_settings-button']"))
            )
            settings_button.click()
            time.sleep(2)
            
            # Select Price option
            try:
                # Open Display Value dropdown
                dropdown = driver.find_element(By.XPATH, "//div[div[text()='Display value']]/following-sibling::div[1]")
                dropdown.click()
                time.sleep(1)
                
                # Use keyboard navigation to select Price
                actions = ActionChains(driver)
                for _ in range(2):
                    actions.send_keys(Keys.ARROW_DOWN).perform()
                    time.sleep(0.5)
                actions.send_keys(Keys.RETURN).perform()
                time.sleep(1)
                
            except Exception:
                pass  # If price selection fails, continue with default view

            # Configure color depth
            #print("Setting color depth...")
            try:
                # Find and click color depth button
                color_button = wait.until(
                    EC.element_to_be_clickable((By.CSS_SELECTOR, "[data-qa-id='heatmap-bottom-bar_button_multiplier']"))
                )
                color_button.click()
               #print("Clicked color depth button")
                time.sleep(1)

                # Navigate using keyboard
                #print("Navigating color depth options...")
                actions = ActionChains(driver)
                # Press down arrow three times
                for _ in range(3):
                    actions.send_keys(Keys.ARROW_DOWN).perform()
                    time.sleep(0.5)
                
                # Press Enter to select
                actions.send_keys(Keys.RETURN).perform()
                #print("Selected color depth option")
                
                # Move mouse away and click in blank area to remove highlight
                #print("Removing button highlight...")
                actions.move_by_offset(200, 0).click().perform()
            except Exception as e:
                #print(f"Warning: Could not set color depth - {str(e)}")
                pass

            # Take screenshot
            output_file = os.path.join("CodeOutput", "stock_heatmap_price.png")
            driver.save_screenshot(output_file)
            
            if os.path.exists(output_file):
                print("Function 4 - NIFTY50 Heatmap - Successful")
                return True
            else:
                print("Function 4 - NIFTY50 Heatmap - Not Successful")
                return False

        finally:
            driver.quit()

    except Exception as e:
        print("Function 4 - NIFTY50 Heatmap - Not Successful")
        return False

def get_nifty_gainers_losers():
    """Get NIFTY50 Top 5 Gainers and Losers"""
    try:
        nifty50_tickers = [
            "ADANIENT.NS", "ADANIPORTS.NS", "APOLLOHOSP.NS", "ASIANPAINT.NS", 
            "AXISBANK.NS", "BAJAJ-AUTO.NS", "BAJFINANCE.NS", "BAJAJFINSV.NS", 
            "BPCL.NS", "BHARTIARTL.NS", "BRITANNIA.NS", "CIPLA.NS", "COALINDIA.NS", 
            "DIVISLAB.NS", "DRREDDY.NS", "EICHERMOT.NS", "GRASIM.NS", "HCLTECH.NS", 
            "HDFCBANK.NS", "HDFCLIFE.NS", "HEROMOTOCO.NS", "HINDALCO.NS", 
            "HINDUNILVR.NS", "ICICIBANK.NS", "ITC.NS", "INDUSINDBK.NS", "INFY.NS", 
            "JSWSTEEL.NS", "KOTAKBANK.NS", "LTIM.NS", "LT.NS", "M&M.NS", 
            "MARUTI.NS", "NTPC.NS", "NESTLEIND.NS", "ONGC.NS", "POWERGRID.NS", 
            "RELIANCE.NS", "SBILIFE.NS", "SBIN.NS", "SUNPHARMA.NS", "TCS.NS", 
            "TATACONSUM.NS", "TATAMOTORS.NS", "TATASTEEL.NS", "TECHM.NS", 
            "TITAN.NS", "ULTRACEMCO.NS", "WIPRO.NS", "SHRIRAMFIN.NS"
        ]

        # Download data for the last 2 trading days
        data = yf.download(nifty50_tickers, period="2d", auto_adjust=True)
        if data.empty:
            return False

        # Get close prices and calculate changes
        close_prices = data['Close']
        if len(close_prices) < 2:
            return False

        latest_prices = close_prices.iloc[-1]
        previous_prices = close_prices.iloc[-2]
        percent_change = ((latest_prices - previous_prices) / previous_prices) * 100
        
        # Process results
        results = []
        for ticker in percent_change.index:
            if pd.notna(percent_change[ticker]) and pd.notna(latest_prices[ticker]):
                results.append({
                    "stock": ticker.replace(".NS", ""),
                    "price": latest_prices[ticker],
                    "change": percent_change[ticker]
                })

        sorted_results = sorted(results, key=lambda x: x['change'], reverse=True)
        top_gainers = sorted_results[:5]
        top_losers = sorted_results[-5:]
        top_losers.reverse()

        # Save results to text file
        output_file = os.path.join("CodeOutput", "nifty50_movers.txt")
        with open(output_file, 'w') as f:
            f.write("NIFTY 50 Top Movers Report\n")
            f.write("=" * 40 + "\n\n")
            
            f.write("Top 5 Gainers:\n")
            f.write("=" * 40 + "\n")
            f.write(f"{'Stock':<15}{'Price':<12}{'Change (%)':<10}\n")
            f.write("-" * 40 + "\n")
            for gainer in top_gainers:
                f.write(f"{gainer['stock']:<15}{gainer['price']:<12.2f}+{gainer['change']:.2f}%\n")
            
            f.write("\nTop 5 Losers:\n")
            f.write("=" * 40 + "\n")
            f.write(f"{'Stock':<15}{'Price':<12}{'Change (%)':<10}\n")
            f.write("-" * 40 + "\n")
            for loser in top_losers:
                f.write(f"{loser['stock']:<15}{loser['price']:<12.2f}{loser['change']:.2f}%\n")

        print("Function 3 - NIFTY50 Top 5 Gainers & Losers - Successful")
        return True

    except Exception as e:
        print("Function 3 - NIFTY50 Top 5 Gainers & Losers - Not Successful")
        return False

def get_nifty_seven_days():
    """Generate NIFTY50 7-day analysis report"""
    try:
        # Download and analyze data
        nifty_daily_data = yf.download("^NSEI", period="90d", interval="1d", auto_adjust=True)
        nifty_hourly_data = yf.download("^NSEI", period="15d", interval="1h", auto_adjust=True)

        # Flatten multi-level columns if present
        if isinstance(nifty_daily_data.columns, pd.MultiIndex):
            nifty_daily_data.columns = nifty_daily_data.columns.get_level_values(0)
        if isinstance(nifty_hourly_data.columns, pd.MultiIndex):
            nifty_hourly_data.columns = nifty_hourly_data.columns.get_level_values(0)

        if nifty_daily_data.empty or nifty_hourly_data.empty:
            return False

        # Analyze data
        nifty_daily_data['SMA50'] = nifty_daily_data['Close'].rolling(window=50).mean()
        latest_close = nifty_daily_data['Close'].iloc[-1]
        latest_sma50 = nifty_daily_data['SMA50'].iloc[-1]
        
        recent_period = nifty_daily_data.tail(15)
        resistance_level = float(recent_period['High'].max())
        support_level = float(recent_period['Low'].min())
        
        analysis_data = {
            "daily_data": nifty_daily_data,
            "hourly_data": nifty_hourly_data,
            "latest_close": latest_close,
            "latest_sma50": latest_sma50,
            "resistance_1": resistance_level,
            "resistance_2": resistance_level * 1.005,
            "support_1": support_level,
            "support_2": support_level * 0.995,
        }

        # Generate chart
        chart_file = os.path.join("CodeOutput", "nifty_1hr_chart.png")
        hourly_df = analysis_data['hourly_data'].tail(7*8).copy()
        
        cols_to_check = ['Open', 'High', 'Low', 'Close']
        for col in cols_to_check:
            hourly_df[col] = pd.to_numeric(hourly_df[col], errors='coerce')
        hourly_df.dropna(inplace=True)

        if hourly_df.empty:
            return False

        hlines = dict(hlines=[analysis_data['resistance_1'], analysis_data['support_1']], 
                     colors=['r', 'g'], linestyle='--')
        mc = mpf.make_marketcolors(up='#00b746', down='#ef403c', inherit=True)
        style = mpf.make_mpf_style(marketcolors=mc, base_mpf_style='nightclouds')

        mpf.plot(
            hourly_df, type='candle', style=style, title='Nifty 50 - 1 Hour Chart',
            ylabel='Price (INR)', volume=False, hlines=hlines, figratio=(16, 9),
            savefig=dict(fname=chart_file, dpi=150, pad_inches=0.1)
        )

        # Generate PDF report
        pdf_file = os.path.join("CodeOutput", "Market_Report_Segment_Nifty50.pdf")
        c = canvas.Canvas(pdf_file, pagesize=letter)
        width, height = letter

        # Add PDF content
        title_color = HexColor("#1E3A8A")
        text_color = HexColor("#1F2937")
        resistance_color = HexColor("#B91C1C")
        support_color = HexColor("#15803D")

        styles = getSampleStyleSheet()
        body_style = styles['BodyText']
        body_style.textColor = text_color
        body_style.fontSize = 11
        body_style.leading = 14

        c.setFillColor(title_color)
        c.setFont("Helvetica-Bold", 18)
        c.drawString(1 * inch, height - 1 * inch, "Nifty Technical Analysis")
        c.line(1*inch, height - 1.05*inch, width - 1*inch, height - 1.05*inch)

        # Add chart to PDF
        if os.path.exists(chart_file):
            chart_width = 6.5 * inch
            chart_height = (chart_width / 16) * 9
            c.drawImage(chart_file, 1 * inch, height - 1.5 * inch - chart_height, 
                       width=chart_width, height=chart_height)
            y_position = height - 1.8 * inch - chart_height
        else:
            c.drawString(1*inch, height - 1.7*inch, "Chart could not be generated.")
            y_position = height - 2.2 * inch

        c.setFillColor(title_color)
        c.setFont("Helvetica-Bold", 14)
        c.drawString(1 * inch, y_position, "◆ Overview")
        y_position -= 0.3 * inch

        daily_analysis_text = f"""
        <b>Daily Chart Analysis:</b> Nifty is currently trading around {analysis_data['latest_close']:.2f}, which is 
        {'above' if analysis_data['latest_close'] > analysis_data['latest_sma50'] else 'below'} the key 50-day SMA of 
        {analysis_data['latest_sma50']:.2f}. This indicates a {'bullish' if analysis_data['latest_close'] > analysis_data['latest_sma50'] else 'bearish'}
        medium-term trend. The index has shown {'strength' if analysis_data['latest_close'] > analysis_data['resistance_1']*0.98 else 'weakness'} 
        in recent sessions.
        """
        hourly_analysis_text = f"""
        <b>1-Hour Chart Analysis:</b> The short-term chart shows the price action consolidating near the recent 
        highs. A breakout above the immediate resistance at {analysis_data['resistance_1']:.2f} could trigger further 
        upside momentum, while a failure to hold the support at {analysis_data['support_1']:.2f} may lead to a pullback.
        """
        
        p = Paragraph(daily_analysis_text, body_style)
        p.wrapOn(c, width - 2*inch, height)
        p.drawOn(c, 1*inch, y_position - p.height)
        y_position -= (p.height + 0.2*inch)

        p = Paragraph(hourly_analysis_text, body_style)
        p.wrapOn(c, width - 2*inch, height)
        p.drawOn(c, 1*inch, y_position - p.height)
        y_position -= (p.height + 0.3*inch)

        c.setFillColor(title_color)
        c.setFont("Helvetica-Bold", 14)
        c.drawString(1 * inch, y_position, "◆ Key Levels to Watch")
        y_position -= 0.3 * inch

        c.setFillColor(resistance_color)
        c.setFont("Helvetica-Bold", 12)
        c.drawString(1.1 * inch, y_position, "RESISTANCE")
        c.setFont("Helvetica", 11)
        y_position -= 0.25 * inch
        c.drawString(1.2 * inch, y_position, f"● {analysis_data['resistance_1']:.2f} – Immediate resistance from recent highs.")
        y_position -= 0.25 * inch
        c.drawString(1.2 * inch, y_position, f"● {analysis_data['resistance_2']:.2f} – Next potential resistance zone.")
        y_position -= 0.4 * inch

        c.setFillColor(support_color)
        c.setFont("Helvetica-Bold", 12)
        c.drawString(1.1 * inch, y_position, "SUPPORT")
        c.setFont("Helvetica", 11)
        y_position -= 0.25 * inch
        c.drawString(1.2 * inch, y_position, f"● {analysis_data['support_1']:.2f} – Immediate support from recent lows.")
        y_position -= 0.25 * inch
        c.drawString(1.2 * inch, y_position, f"● {analysis_data['support_2']:.2f} – Next strong support zone.")

        c.save()

        # Clean up
        if os.path.exists(chart_file):
            os.remove(chart_file)

        print("Function 2 - NIFTY50 Last 7 days analysis - Successful")
        return True

    except Exception as e:
        print("Function 2 - NIFTY50 Last 7 days analysis - Not Successful")
        return False


if __name__ == "__main__":
    # Create output directory
    os.makedirs("CodeOutput", exist_ok=True)
    
    # Initialize success counter
    success_count = 0
    
    # Execute all functions in order and track successes
    if get_nifty_summary():                  # 1
        success_count += 1
    if get_nifty_seven_days():              # 2
        success_count += 1
    if get_nifty_gainers_losers():          # 3
        success_count += 1
    if get_nifty_heatmap():                 # 4
        success_count += 1
    if get_nifty_pcr():                     # 5
        success_count += 1
    if get_nifty_oi():                      # 6
        success_count += 1
    if get_market_news():                   # 7
        success_count += 1
    if get_key_stocks_to_watch():           # 8
        success_count += 1
    if get_vix_analysis():                  # 9
        success_count += 1
    if get_fii_dii_data():                 # 10
        success_count += 1
    if get_sgx_nifty():                    # 11
        success_count += 1
    if get_global_markets():               # 12
        success_count += 1
    if get_gold_rates():                   # 13
        success_count += 1
    if get_silver_rates():                 # 14
        success_count += 1
    if get_currency_rates():               # 15
        success_count += 1
        
    # Print final success count
    print(f"\n{success_count} / 15 functions executed successfully")
